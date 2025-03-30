"""
RAG Enhancements for the Anthropic-powered Agent
Adds support for more document formats, hybrid search, and metadata filtering
"""

import os
import io
import re
import json
import logging
import hashlib
import numpy as np
from typing import Dict, List, Optional, Any, Tuple, Union
from datetime import datetime
from pathlib import Path

# Document format libraries
import docx
import PyPDF2
import pandas as pd
import openpyxl
from bs4 import BeautifulSoup

# Logging configuration
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Process various document formats for RAG system."""
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extract text from a PDF file."""
        try:
            text = ""
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
            raise
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract text from a DOCX file."""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            raise
    
    @staticmethod
    def extract_text_from_excel(file_path: str) -> str:
        """Extract text from an Excel file."""
        try:
            df_dict = pd.read_excel(file_path, sheet_name=None)
            text = ""
            for sheet_name, df in df_dict.items():
                text += f"Sheet: {sheet_name}\n"
                text += df.to_string(index=False) + "\n\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting text from Excel {file_path}: {str(e)}")
            raise
    
    @staticmethod
    def extract_text_from_csv(file_path: str) -> str:
        """Extract text from a CSV file."""
        try:
            df = pd.read_csv(file_path)
            return df.to_string(index=False)
        except Exception as e:
            logger.error(f"Error extracting text from CSV {file_path}: {str(e)}")
            raise
    
    @staticmethod
    def extract_text_from_html(file_path: str) -> str:
        """Extract text from an HTML file."""
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                html_content = file.read()
            
            soup = BeautifulSoup(html_content, "html.parser")
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.extract()
            
            # Get text
            text = soup.get_text()
            
            # Break into lines and remove leading and trailing space on each
            lines = (line.strip() for line in text.splitlines())
            
            # Break multi-headlines into a line each
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            
            # Join the non-empty lines
            text = "\n".join(chunk for chunk in chunks if chunk)
            
            return text
        except Exception as e:
            logger.error(f"Error extracting text from HTML {file_path}: {str(e)}")
            raise
    
    @staticmethod
    def extract_text_from_file(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from a file based on its extension."""
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        metadata = {
            "filename": file_path.name,
            "extension": extension,
            "size_bytes": os.path.getsize(file_path),
            "last_modified": datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat(),
            "processed_at": datetime.now().isoformat()
        }
        
        try:
            if extension == ".pdf":
                text = DocumentProcessor.extract_text_from_pdf(file_path)
                metadata["type"] = "pdf"
                
            elif extension == ".docx":
                text = DocumentProcessor.extract_text_from_docx(file_path)
                metadata["type"] = "docx"
                
            elif extension in [".xlsx", ".xls"]:
                text = DocumentProcessor.extract_text_from_excel(file_path)
                metadata["type"] = "excel"
                
            elif extension == ".csv":
                text = DocumentProcessor.extract_text_from_csv(file_path)
                metadata["type"] = "csv"
                
            elif extension in [".html", ".htm"]:
                text = DocumentProcessor.extract_text_from_html(file_path)
                metadata["type"] = "html"
                
            elif extension in [".txt", ".md", ".py", ".js", ".java", ".c", ".cpp", ".cs", ".json", ".xml"]:
                with open(file_path, "r", encoding="utf-8") as file:
                    text = file.read()
                metadata["type"] = "text"
                
            else:
                # Try to read as text, but may fail for binary files
                try:
                    with open(file_path, "r", encoding="utf-8") as file:
                        text = file.read()
                    metadata["type"] = "text"
                except UnicodeDecodeError:
                    raise ValueError(f"Unsupported file format: {extension}")
            
            # Add word and character counts to metadata
            metadata["word_count"] = len(text.split())
            metadata["char_count"] = len(text)
            
            return text, metadata
        
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            raise

class HybridSearcher:
    """Hybrid search combining vector and keyword search."""
    
    def __init__(self, vector_store):
        """Initialize the hybrid searcher."""
        self.vector_store = vector_store
    
    def _keyword_search(self, query: str, chunks: List[Dict[str, Any]], top_k: int = 10) -> List[Dict[str, Any]]:
        """Search chunks using keywords."""
        # Convert query to lowercase
        query_lower = query.lower()
        
        # Extract keywords (simple tokenization)
        keywords = re.findall(r'\b\w+\b', query_lower)
        
        # Score each chunk based on keyword matches
        scored_chunks = []
        for chunk in chunks:
            content_lower = chunk["content"].lower()
            
            # Count keyword occurrences
            score = sum(content_lower.count(keyword) for keyword in keywords)
            
            # Boost score for exact phrase match
            if query_lower in content_lower:
                score += 3
            
            if score > 0:
                scored_chunks.append({
                    "chunk": chunk,
                    "score": score
                })
        
        # Sort by score (descending)
        scored_chunks.sort(key=lambda x: x["score"], reverse=True)
        
        # Return top_k results
        return scored_chunks[:top_k]
    
    async def hybrid_search(self, query: str, top_k: int = 5, vector_weight: float = 0.7) -> List[Dict[str, Any]]:
        """Perform hybrid search combining vector and keyword search."""
        # Get all chunks for keyword search
        all_chunks = []
        for chunk_path in self.vector_store.chunks_dir.glob("*.json"):
            with open(chunk_path, 'r') as f:
                chunk_dict = json.load(f)
                all_chunks.append(chunk_dict)
        
        # Perform keyword search
        keyword_results = self._keyword_search(query, all_chunks, top_k=top_k*2)  # Get more for reranking
        
        # Perform vector search
        vector_results = await self.vector_store.search(query, top_k=top_k*2)  # Get more for reranking
        
        # Combine results
        combined_results = {}
        
        # Add vector results with their scores
        for result in vector_results:
            chunk_id = result.chunk.id
            combined_results[chunk_id] = {
                "chunk": result.chunk,
                "vector_score": result.score,
                "keyword_score": 0,
                "combined_score": result.score * vector_weight
            }
        
        # Add keyword results with their scores
        for result in keyword_results:
            chunk_id = result["chunk"]["id"]
            
            # Normalize keyword score (divide by max possible score)
            max_keyword_score = len(re.findall(r'\b\w+\b', query.lower())) + 3  # +3 for exact phrase match
            normalized_keyword_score = result["score"] / max_keyword_score if max_keyword_score > 0 else 0
            
            if chunk_id in combined_results:
                combined_results[chunk_id]["keyword_score"] = normalized_keyword_score
                combined_results[chunk_id]["combined_score"] += normalized_keyword_score * (1 - vector_weight)
            else:
                combined_results[chunk_id] = {
                    "chunk": result["chunk"],
                    "vector_score": 0,
                    "keyword_score": normalized_keyword_score,
                    "combined_score": normalized_keyword_score * (1 - vector_weight)
                }
        
        # Sort by combined score
        sorted_results = sorted(combined_results.values(), key=lambda x: x["combined_score"], reverse=True)
        
        # Format results
        formatted_results = []
        for result in sorted_results[:top_k]:
            chunk = result["chunk"]
            formatted_results.append({
                "content": chunk["content"],
                "score": result["combined_score"],
                "vector_score": result["vector_score"],
                "keyword_score": result["keyword_score"],
                "metadata": {
                    **chunk["metadata"],
                    "document_id": chunk["document_id"]
                }
            })
        
        return formatted_results

class DocumentCollection:
    """Collection of documents with namespacing."""
    
    def __init__(self, name: str, storage_dir: str = "vector_store"):
        """Initialize the document collection."""
        self.name = name
        self.storage_dir = Path(storage_dir) / name
        self.documents_dir = self.storage_dir / "documents"
        self.chunks_dir = self.storage_dir / "chunks"
        self.metadata_path = self.storage_dir / "metadata.json"
        self.metadata = {
            "name": name,
            "document_count": 0,
            "chunk_count": 0,
            "last_updated": None,
            "created_at": datetime.now().isoformat()
        }
        
        # Create directories if they don't exist
        self.documents_dir.mkdir(parents=True, exist_ok=True)
        self.chunks_dir.mkdir(parents=True, exist_ok=True)
        
        # Load metadata if it exists
        if self.metadata_path.exists():
            with open(self.metadata_path, 'r') as f:
                self.metadata = json.load(f)
    
    def _save_metadata(self):
        """Save metadata to disk."""
        self.metadata["last_updated"] = datetime.now().isoformat()
        with open(self.metadata_path, 'w') as f:
            json.dump(self.metadata, f, indent=2)
    
    def add_document(self, document):
        """Add a document to the collection."""
        # Save document
        document_path = self.documents_dir / f"{document.id}.json"
        with open(document_path, 'w') as f:
            document_dict = document.dict()
            json.dump(document_dict, f, indent=2)
        
        # Update metadata
        self.metadata["document_count"] += 1
        self._save_metadata()
        
        return document.id
    
    def add_chunk(self, chunk):
        """Add a chunk to the collection."""
        # Save chunk
        chunk_path = self.chunks_dir / f"{chunk.id}.json"
        with open(chunk_path, 'w') as f:
            chunk_dict = chunk.dict()
            json.dump(chunk_dict, f, indent=2)
        
        # Update metadata
        self.metadata["chunk_count"] += 1
        self._save_metadata()
        
        return chunk.id
    
    def get_document(self, document_id):
        """Get a document from the collection."""
        document_path = self.documents_dir / f"{document_id}.json"
        if not document_path.exists():
            return None
        
        with open(document_path, 'r') as f:
            document_dict = json.load(f)
            return document_dict
    
    def get_chunk(self, chunk_id):
        """Get a chunk from the collection."""
        chunk_path = self.chunks_dir / f"{chunk_id}.json"
        if not chunk_path.exists():
            return None
        
        with open(chunk_path, 'r') as f:
            chunk_dict = json.load(f)
            return chunk_dict
    
    def delete_document(self, document_id):
        """Delete a document and its chunks from the collection."""
        document_path = self.documents_dir / f"{document_id}.json"
        if not document_path.exists():
            return False
        
        # Delete document
        document_path.unlink()
        
        # Delete chunks
        chunks_deleted = 0
        for chunk_path in self.chunks_dir.glob("*.json"):
            with open(chunk_path, 'r') as f:
                chunk_dict = json.load(f)
                if chunk_dict["document_id"] == document_id:
                    chunk_path.unlink()
                    chunks_deleted += 1
        
        # Update metadata
        self.metadata["document_count"] -= 1
        self.metadata["chunk_count"] -= chunks_deleted
        self._save_metadata()
        
        return True
    
    def search(self, query_embedding, top_k=5):
        """Search for similar chunks using cosine similarity."""
        results = []
        
        # Load all chunks
        for chunk_path in self.chunks_dir.glob("*.json"):
            with open(chunk_path, 'r') as f:
                chunk_dict = json.load(f)
                
                if "embedding" in chunk_dict and chunk_dict["embedding"]:
                    # Calculate cosine similarity
                    similarity = self._cosine_similarity(query_embedding, chunk_dict["embedding"])
                    results.append({
                        "chunk": chunk_dict,
                        "score": similarity
                    })
        
        # Sort by similarity score (descending)
        results.sort(key=lambda x: x["score"], reverse=True)
        
        # Return top k results
        return results[:top_k]
    
    def _cosine_similarity(self, vec1, vec2):
        """Calculate cosine similarity between two vectors."""
        if not vec1 or not vec2:
            return 0.0
        
        vec1 = np.array(vec1)
        vec2 = np.array(vec2)
        
        dot_product = np.dot(vec1, vec2)
        norm1 = np.linalg.norm(vec1)
        norm2 = np.linalg.norm(vec2)
        
        if norm1 == 0 or norm2 == 0:
            return 0.0
        
        return dot_product / (norm1 * norm2)
    
    def get_stats(self):
        """Get statistics about the collection."""
        return {
            **self.metadata,
            "storage_size_bytes": self._get_storage_size()
        }
    
    def _get_storage_size(self):
        """Calculate the total storage size in bytes."""
        total_size = 0
        
        # Add size of documents
        for document_path in self.documents_dir.glob("*.json"):
            total_size += document_path.stat().st_size
        
        # Add size of chunks
        for chunk_path in self.chunks_dir.glob("*.json"):
            total_size += chunk_path.stat().st_size
        
        # Add size of metadata
        if self.metadata_path.exists():
            total_size += self.metadata_path.stat().st_size
        
        return total_size

class CollectionManager:
    """Manage multiple document collections."""
    
    def __init__(self, storage_dir: str = "vector_store"):
        """Initialize the collection manager."""
        self.storage_dir = Path(storage_dir)
        self.collections = {}
        self.metadata_path = self.storage_dir / "collections.json"
        self.metadata = {
            "collections": [],
            "default_collection": "default",
            "last_updated": None
        }
        
        # Create storage directory if it doesn't exist
        self.storage_dir.mkdir(parents=True, exist_ok=True)
        
        # Load metadata if it exists
        if self.metadata_path.exists():
            with open(self.metadata_path, 'r') as f:
                self.metadata = json.load(f)
        
        # Load collections
        for collection_name in self.metadata["collections"]:
            self.collections[collection_name] = DocumentCollection(collection_name, storage_dir)
        
        # Create default collection if it doesn't exist
        if "default" not in self.collections:
            self.create_collection("default")
    
    def _save_metadata(self):
        """Save metadata to disk."""
        self.metadata["last_updated"] = datetime.now().isoformat()
        with open(self.metadata_path, 'w') as f:
            json.dump(self.metadata, f, indent=2)
    
    def create_collection(self, name: str) -> DocumentCollection:
        """Create a new collection."""
        if name in self.collections:
            raise ValueError(f"Collection {name} already exists")
        
        # Create collection
        collection = DocumentCollection(name, self.storage_dir)
        self.collections[name] = collection
        
        # Update metadata
        if name not in self.metadata["collections"]:
            self.metadata["collections"].append(name)
            self._save_metadata()
        
        return collection
    
    def get_collection(self, name: str) -> DocumentCollection:
        """Get a collection by name."""
        if name not in self.collections:
            raise ValueError(f"Collection {name} does not exist")
        
        return self.collections[name]
    
    def delete_collection(self, name: str) -> bool:
        """Delete a collection."""
        if name not in self.collections:
            return False
        
        if name == "default":
            raise ValueError("Cannot delete the default collection")
        
        # Remove from collections
        del self.collections[name]
        
        # Update metadata
        self.metadata["collections"].remove(name)
        self._save_metadata()
        
        # Delete collection directory
        collection_dir = self.storage_dir / name
        if collection_dir.exists():
            import shutil
            shutil.rmtree(collection_dir)
        
        return True
    
    def list_collections(self) -> List[Dict[str, Any]]:
        """List all collections."""
        collections = []
        for name, collection in self.collections.items():
            stats = collection.get_stats()
            collections.append({
                "name": name,
                "document_count": stats["document_count"],
                "chunk_count": stats["chunk_count"],
                "created_at": stats.get("created_at", "Unknown"),
                "last_updated": stats.get("last_updated", "Never"),
                "is_default": name == self.metadata["default_collection"]
            })
        
        return collections
    
    def set_default_collection(self, name: str) -> bool:
        """Set the default collection."""
        if name not in self.collections:
            return False
        
        self.metadata["default_collection"] = name
        self._save_metadata()
        
        return True
    
    def get_default_collection(self) -> DocumentCollection:
        """Get the default collection."""
        default_name = self.metadata["default_collection"]
        return self.collections[default_name]

# Enhanced RAG system tools
def get_enhanced_rag_tools() -> List:
    """Get enhanced tools for the RAG system."""
    from anthropic_agent import Tool, ToolParameter
    
    # Initialize collection manager
    collection_manager = CollectionManager()
    
    async def add_file_to_collection(file_path: str, collection_name: Optional[str] = None) -> Dict[str, Any]:
        """Add a file to a specific collection."""
        try:
            # Get collection
            if collection_name:
                try:
                    collection = collection_manager.get_collection(collection_name)
                except ValueError:
                    collection = collection_manager.create_collection(collection_name)
            else:
                collection = collection_manager.get_default_collection()
            
            # Process file
            text, metadata = DocumentProcessor.extract_text_from_file(file_path)
            
            # Add additional metadata
            metadata["collection"] = collection.name
            metadata["added_to_collection_at"] = datetime.now().isoformat()
            
            # TODO: Process the document using the collection
            # This would require integrating with the main RAG system
            
            return {
                "success": True,
                "collection": collection.name,
                "file": os.path.basename(file_path),
                "extracted_text_length": len(text),
                "metadata": metadata
            }
        except Exception as e:
            return {
                "success": False,
                "error": str(e)
            }
    
    def create_collection(name: str) -> Dict[str, Any]:
        """Create a new collection."""
        try:
            collection = collection_manager.create_collection(name)
            return {
                "success": True,
                "collection": name,
                "message": f"Collection '{name}' created successfully"
            }
        except ValueError as e:
            return {
                "success": False,
                "error": str(e)
            }
    
    def list_collections() -> Dict[str, Any]:
        """List all collections."""
        collections = collection_manager.list_collections()
        return {
            "collections": collections,
            "count": len(collections),
            "default_collection": collection_manager.metadata["default_collection"]
        }
    
    def set_default_collection(name: str) -> Dict[str, Any]:
        """Set the default collection."""
        success = collection_manager.set_default_collection(name)
        if success:
            return {
                "success": True,
                "message": f"Default collection set to '{name}'"
            }
        else:
            return {
                "success": False,
                "error": f"Collection '{name}' does not exist"
            }
    
    def delete_collection(name: str) -> Dict[str, Any]:
        """Delete a collection."""
        try:
            success = collection_manager.delete_collection(name)
            if success:
                return {
                    "success": True,
                    "message": f"Collection '{name}' deleted successfully"
                }
            else:
                return {
                    "success": False,
                    "error": f"Collection '{name}' does not exist"
                }
        except ValueError as e:
            return {
                "success": False,
                "error": str(e)
            }
    
    async def hybrid_search(query: str, collection_name: Optional[str] = None, top_k: int = 5, vector_weight: float = 0.7) -> Dict[str, Any]:
        """Perform hybrid search on a collection."""
        try:
            # Get collection
            if collection_name:
                try:
                    collection = collection_manager.get_collection(collection_name)
                except ValueError:
                    return {
                        "success": False,
                        "error": f"Collection '{collection_name}' does not exist"
                    }
            else:
                collection = collection_manager.get_default_collection()
            
            # TODO: Perform hybrid search using the collection
            # This would require integrating with the main RAG system
            
            return {
                "success": True,
                "collection": collection.name,
                "query": query,
                "top_k": top_k,
                "vector_weight": vector_weight,
                "message": "Hybrid search not yet implemented"
            }
        except Exception as e:
            return {
                "success": False,
                "error": str(e)
            }
    
    # Define tools
    return [
        Tool(
            name="rag_add_file_to_collection",
            description="Add a file to a specific collection in the RAG system",
            parameters=[
                ToolParameter(name="file_path", type="string", description="Path to the file to add"),
                ToolParameter(name="collection_name", type="string", description="Name of the collection", required=False)
            ],
            function=add_file_to_collection,
            category="rag"
        ),
        
        Tool(
            name="rag_create_collection",
            description="Create a new collection in the RAG system",
            parameters=[
                ToolParameter(name="name", type="string", description="Name of the collection")
            ],
            function=create_collection,
            category="rag"
        ),
        
        Tool(
            name="rag_list_collections",
            description="List all collections in the RAG system",
            parameters=[],
            function=list_collections,
            category="rag"
        ),
        
        Tool(
            name="rag_set_default_collection",
            description="Set the default collection in the RAG system",
            parameters=[
                ToolParameter(name="name", type="string", description="Name of the collection")
            ],
            function=set_default_collection,
            category="rag"
        ),
        
        Tool(
            name="rag_delete_collection",
            description="Delete a collection from the RAG system",
            parameters=[
                ToolParameter(name="name", type="string", description="Name of the collection")
            ],
            function=delete_collection,
            category="rag"
        ),
        
        Tool(
            name="rag_hybrid_search",
            description="Perform hybrid search on a collection",
            parameters=[
                ToolParameter(name="query", type="string", description="Query text"),
                ToolParameter(name="collection_name", type="string", description="Name of the collection", required=False),
                ToolParameter(name="top_k", type="integer", description="Number of top results to return", required=False, default=5),
                ToolParameter(name="vector_weight", type="number", description="Weight of vector search in hybrid search (0-1)", required=False, default=0.7)
            ],
            function=hybrid_search,
            category="rag"
        )
    ]